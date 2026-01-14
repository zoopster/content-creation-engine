"""
DOCX Generation Skill - Creates Microsoft Word documents from draft content.

This skill transforms markdown content into formatted DOCX files with:
- Professional styling and formatting
- Template support
- Headers, footers, and page numbers
- Tables, lists, and rich text formatting
"""

import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), '../..'))

from typing import Dict, Any, Optional
from agents.base.agent import Skill
from agents.base.models import DraftContent
from datetime import datetime
import re


class DocxGenerationSkill(Skill):
    """
    Generates Microsoft Word (DOCX) documents from draft content.

    Features:
    - Markdown to DOCX conversion
    - Template application
    - Style and formatting preservation
    - Header and footer support
    """

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__("docx-generation", config)
        self.output_dir = config.get("output_dir", "./output") if config else "./output"
        self.template_dir = config.get("template_dir", "./templates") if config else "./templates"

        # Ensure directories exist
        os.makedirs(self.output_dir, exist_ok=True)
        os.makedirs(self.template_dir, exist_ok=True)

    def execute(self, input_data: DraftContent, **kwargs) -> Dict[str, Any]:
        """
        Generate DOCX document from draft content.

        Args:
            input_data: DraftContent object with markdown content
            **kwargs:
                - brand_template: BrandTemplate for styling (Phase 3)
                - template_name: Optional template to use
                - include_metadata: Include metadata page (default: True)
                - page_numbers: Add page numbers (default: True)
                - header_text: Custom header text
                - footer_text: Custom footer text

        Returns:
            Dictionary with:
                - file_path: Path to generated DOCX file
                - success: Boolean indicating success
                - metadata: Additional information
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            self.logger.error("python-docx not installed. Install with: pip install python-docx")
            # Return mock result for development
            return self._generate_mock_docx(input_data, **kwargs)

        brand_template = kwargs.get("brand_template")
        template_name = kwargs.get("template_name")
        include_metadata = kwargs.get("include_metadata", True)

        # Load template or create new document
        doc = self._load_template(template_name) if template_name else Document()

        # Apply brand styling if template provided
        if brand_template:
            self._apply_brand_styles(doc, brand_template)
            self._set_page_layout(doc, brand_template)

        # Add metadata page if requested
        if include_metadata:
            self._add_metadata_page(doc, input_data, brand_template)
            doc.add_page_break()

        # Parse and add content
        self._parse_markdown_to_docx(doc, input_data.content, brand_template)

        # Add headers/footers
        header_text = kwargs.get("header_text")
        footer_text = kwargs.get("footer_text")
        if brand_template:
            # Use brand template for header/footer
            header_text = header_text or brand_template.company_name
            footer_text = footer_text or f"© {datetime.now().year} {brand_template.company_name}"

        if header_text or footer_text or kwargs.get("page_numbers", True):
            self._add_headers_footers(doc, header_text, footer_text, kwargs.get("page_numbers", True))

        # Generate filename and save
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{input_data.content_type.value}_{timestamp}.docx"
        file_path = os.path.join(self.output_dir, filename)

        doc.save(file_path)
        self.logger.info(f"Generated DOCX: {file_path}")

        return {
            "file_path": file_path,
            "success": True,
            "metadata": {
                "word_count": input_data.word_count,
                "pages": len(doc.sections),
                "template": template_name,
                "created_at": datetime.now().isoformat()
            }
        }

    def _apply_brand_styles(self, doc: Any, brand_template: Any):
        """Apply brand template styles to document."""
        from docx.shared import Pt, RGBColor

        try:
            # Update Normal style
            style = doc.styles['Normal']
            font = style.font
            font.name = brand_template.typography.body_font
            font.size = Pt(brand_template.typography.body_size)

            # Helper to convert hex to RGB
            def hex_to_rgb(hex_color: str) -> tuple:
                """Convert hex color to RGB tuple."""
                hex_color = hex_color.lstrip('#')
                return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Update Heading 1 style
            if 'Heading 1' in doc.styles:
                style = doc.styles['Heading 1']
                style.font.name = brand_template.typography.heading_font
                style.font.size = Pt(brand_template.typography.h1_size)
                r, g, b = hex_to_rgb(brand_template.colors.primary)
                style.font.color.rgb = RGBColor(r, g, b)

            # Update Heading 2 style
            if 'Heading 2' in doc.styles:
                style = doc.styles['Heading 2']
                style.font.name = brand_template.typography.heading_font
                style.font.size = Pt(brand_template.typography.h2_size)
                r, g, b = hex_to_rgb(brand_template.colors.secondary)
                style.font.color.rgb = RGBColor(r, g, b)

            # Update Heading 3 style
            if 'Heading 3' in doc.styles:
                style = doc.styles['Heading 3']
                style.font.name = brand_template.typography.heading_font
                style.font.size = Pt(brand_template.typography.h3_size)
                r, g, b = hex_to_rgb(brand_template.colors.secondary)
                style.font.color.rgb = RGBColor(r, g, b)

            self.logger.info(f"Applied {brand_template.name} brand styles to document")

        except Exception as e:
            self.logger.warning(f"Could not apply all brand styles: {e}")

    def _set_page_layout(self, doc: Any, brand_template: Any):
        """Set page layout from brand template."""
        from docx.shared import Inches

        try:
            section = doc.sections[0]
            layout = brand_template.document_layout

            # Set page dimensions
            section.page_height = Inches(layout.page_height)
            section.page_width = Inches(layout.page_width)

            # Set margins
            section.top_margin = Inches(layout.margin_top)
            section.bottom_margin = Inches(layout.margin_bottom)
            section.left_margin = Inches(layout.margin_left)
            section.right_margin = Inches(layout.margin_right)

            # Set header/footer margins
            section.header_distance = Inches(layout.header_margin)
            section.footer_distance = Inches(layout.footer_margin)

            self.logger.info(f"Applied document layout: {layout.page_size}")

        except Exception as e:
            self.logger.warning(f"Could not apply page layout: {e}")

    def _load_template(self, template_name: str) -> Any:
        """Load a DOCX template."""
        try:
            from docx import Document
            template_path = os.path.join(self.template_dir, f"{template_name}.docx")
            if os.path.exists(template_path):
                return Document(template_path)
        except Exception as e:
            self.logger.warning(f"Could not load template {template_name}: {e}")

        # Return blank document if template not found
        from docx import Document
        return Document()

    def _add_metadata_page(self, doc: Any, draft: DraftContent, brand_template: Any = None):
        """Add a metadata/title page to the document."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(draft.content_type.value.replace("_", " ").title())
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacing

        # Add brand info if provided
        if brand_template:
            brand_para = doc.add_paragraph()
            brand_run = brand_para.add_run(brand_template.company_name)
            brand_run.font.size = Pt(14)
            brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()  # Spacing

        # Metadata table
        metadata_items = [
            ("Word Count", str(draft.word_count)),
            ("Created", datetime.now().strftime("%Y-%m-%d %H:%M")),
            ("Format", draft.format)
        ]

        # Add metadata from draft
        if draft.metadata:
            if "target_audience" in draft.metadata:
                metadata_items.append(("Target Audience", draft.metadata["target_audience"]))
            if "tone" in draft.metadata:
                metadata_items.append(("Tone", draft.metadata["tone"]))

        for label, value in metadata_items:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(value)

    def _parse_markdown_to_docx(self, doc: Any, markdown_content: str, brand_template: Any = None):
        """Parse markdown content and add to DOCX document."""
        lines = markdown_content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Heading 1
            if line.startswith('# '):
                text = line[2:].strip()
                para = doc.add_heading(text, level=1)

            # Heading 2
            elif line.startswith('## '):
                text = line[3:].strip()
                para = doc.add_heading(text, level=2)

            # Heading 3
            elif line.startswith('### '):
                text = line[4:].strip()
                para = doc.add_heading(text, level=3)

            # Blockquote
            elif line.startswith('> '):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='Quote')

            # Unordered list
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                para = doc.add_paragraph(text, style='List Bullet')

            # Ordered list
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                para = doc.add_paragraph(text, style='List Number')

            # Horizontal rule
            elif line.strip() == '---':
                doc.add_paragraph('_' * 50)

            # Code block
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1

                if code_lines:
                    code_text = '\n'.join(code_lines)
                    para = doc.add_paragraph(code_text)
                    para.style = 'No Spacing'
                    # Set monospace font
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)

            # Regular paragraph
            else:
                # Handle inline formatting
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)

            i += 1

    def _add_formatted_text(self, para: Any, text: str):
        """Add text with inline formatting (bold, italic, code)."""
        # Simple regex-based inline formatting
        # This is a simplified version - production would use a proper markdown parser

        parts = []
        current = ""
        i = 0

        while i < len(text):
            # Bold
            if text[i:i+2] == '**':
                if current:
                    parts.append(('normal', current))
                    current = ""
                # Find closing
                end = text.find('**', i+2)
                if end != -1:
                    parts.append(('bold', text[i+2:end]))
                    i = end + 2
                    continue

            # Italic
            elif text[i] == '*' and (i == 0 or text[i-1] != '*'):
                if current:
                    parts.append(('normal', current))
                    current = ""
                # Find closing
                end = text.find('*', i+1)
                if end != -1 and (end+1 >= len(text) or text[end+1] != '*'):
                    parts.append(('italic', text[i+1:end]))
                    i = end + 1
                    continue

            # Inline code
            elif text[i] == '`':
                if current:
                    parts.append(('normal', current))
                    current = ""
                # Find closing
                end = text.find('`', i+1)
                if end != -1:
                    parts.append(('code', text[i+1:end]))
                    i = end + 1
                    continue

            current += text[i]
            i += 1

        if current:
            parts.append(('normal', current))

        # Add formatted runs
        for fmt, txt in parts:
            run = para.add_run(txt)
            if fmt == 'bold':
                run.bold = True
            elif fmt == 'italic':
                run.italic = True
            elif fmt == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

    def _add_headers_footers(self, doc: Any, header_text: Optional[str],
                            footer_text: Optional[str], page_numbers: bool):
        """Add headers and footers to document."""
        section = doc.sections[0]

        # Header
        if header_text:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Footer
        footer = section.footer
        footer_para = footer.paragraphs[0]

        if footer_text:
            footer_para.text = footer_text

        if page_numbers:
            # Add page numbers
            if footer_text:
                footer_para.add_run(" | ")
            footer_para.add_run("Page ")
            # Note: Actual page number fields require XML manipulation
            # This is a simplified version

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _generate_mock_docx(self, draft: DraftContent, **kwargs) -> Dict[str, Any]:
        """
        Generate a mock DOCX file when python-docx is not available.
        Creates a simple text file with .docx extension for development.
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{draft.content_type.value}_{timestamp}.docx"
        file_path = os.path.join(self.output_dir, filename)

        # Create a simple text representation
        content = f"""DOCX Document (Mock)
{'=' * 50}

Content Type: {draft.content_type.value}
Word Count: {draft.word_count}
Created: {datetime.now().isoformat()}

{'=' * 50}

{draft.content}

{'=' * 50}
Note: This is a mock DOCX file. Install python-docx for real DOCX generation.
"""

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

        self.logger.info(f"Generated mock DOCX: {file_path}")

        return {
            "file_path": file_path,
            "success": True,
            "metadata": {
                "word_count": draft.word_count,
                "mock": True,
                "created_at": datetime.now().isoformat()
            }
        }

    def validate_requirements(self) -> tuple[bool, list[str]]:
        """Validate that required dependencies are available."""
        missing = []

        try:
            import docx
        except ImportError:
            missing.append("python-docx")

        return len(missing) == 0, missing
